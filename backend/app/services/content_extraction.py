import base64
import logging
from typing import Dict, Optional, Tuple

logger = logging.getLogger(__name__)


class ContentExtractionService:
    """Service to extract content from various file types for LLM processing."""

    def __init__(self):
        self.max_text_length = 10000  # Limit text length to avoid token limits

    async def extract_content(
        self, file_data: bytes, content_type: str, filename: str
    ) -> Tuple[Optional[str], Dict]:
        """
        Extract content from file data based on content type.

        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            if content_type.startswith("image/"):
                return await self._extract_image_content(file_data, content_type)
            elif content_type == "text/plain":
                return await self._extract_text_content(file_data)
            elif content_type == "application/pdf":
                return await self._extract_pdf_content(file_data)
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ]:
                return await self._extract_word_content(file_data, content_type)
            elif content_type == "application/json":
                return await self._extract_json_content(file_data)
            else:
                # For unsupported types, just return basic info
                return None, {"type": "unsupported", "content_type": content_type}

        except Exception as e:
            logger.error(f"Error extracting content from {filename}: {str(e)}")
            return None, {"error": str(e)}

    async def _extract_image_content(
        self, file_data: bytes, content_type: str
    ) -> Tuple[str, Dict]:
        """Extract image content as base64 for vision models."""
        try:
            # Convert to base64 for LLM vision models
            base64_image = base64.b64encode(file_data).decode("utf-8")

            # Create a description that can be sent to vision-capable models
            image_description = (
                f"[IMAGE: Base64 encoded {content_type} image - {len(file_data)} bytes]"
            )

            metadata = {
                "type": "image",
                "base64_data": base64_image,
                "content_type": content_type,
                "size": len(file_data),
            }

            return image_description, metadata

        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
            return "[IMAGE: Could not process image]", {"error": str(e)}

    async def _extract_text_content(self, file_data: bytes) -> Tuple[str, Dict]:
        """Extract text from plain text files."""
        try:
            # Try different encodings
            for encoding in ["utf-8", "utf-16", "latin-1"]:
                try:
                    text = file_data.decode(encoding)
                    # Limit text length
                    if len(text) > self.max_text_length:
                        text = text[: self.max_text_length] + "... [truncated]"

                    return text, {
                        "type": "text",
                        "encoding": encoding,
                        "length": len(text),
                        "word_count": len(text.split()),
                    }
                except UnicodeDecodeError:
                    continue

            return "[TEXT: Could not decode text file]", {"error": "encoding_error"}

        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            return "[TEXT: Could not process text file]", {"error": str(e)}

    async def _extract_pdf_content(self, file_data: bytes) -> Tuple[str, Dict]:
        """Extract text from PDF files."""
        try:
            # Try to import PyPDF2 or pdfplumber
            try:
                import io

                import PyPDF2

                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
                text_content = []

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(f"Page {page_num + 1}:\n{text}")
                    except Exception as e:
                        text_content.append(
                            f"Page {page_num + 1}: [Could not extract text: {str(e)}]"
                        )

                full_text = "\n\n".join(text_content)

                # Limit text length
                if len(full_text) > self.max_text_length:
                    full_text = full_text[: self.max_text_length] + "... [truncated]"

                return full_text, {
                    "type": "pdf",
                    "page_count": len(pdf_reader.pages),
                    "text_length": len(full_text),
                    "word_count": len(full_text.split()),
                }

            except ImportError:
                return (
                    "[PDF: PyPDF2 not installed - install with: pip install PyPDF2]",
                    {"error": "missing_dependency"},
                )

        except Exception as e:
            logger.error(f"Error extracting PDF: {str(e)}")
            return "[PDF: Could not process PDF file]", {"error": str(e)}

    async def _extract_word_content(
        self, file_data: bytes, content_type: str
    ) -> Tuple[str, Dict]:
        """Extract text from Word documents."""
        try:
            # Try to import python-docx
            try:
                import io

                from docx import Document

                doc = Document(io.BytesIO(file_data))
                text_content = []

                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)

                full_text = "\n".join(text_content)

                # Limit text length
                if len(full_text) > self.max_text_length:
                    full_text = full_text[: self.max_text_length] + "... [truncated]"

                return full_text, {
                    "type": "word_document",
                    "paragraph_count": len(doc.paragraphs),
                    "text_length": len(full_text),
                    "word_count": len(full_text.split()),
                }

            except ImportError:
                return (
                    "[WORD: python-docx not installed - install with: pip install python-docx]",
                    {"error": "missing_dependency"},
                )

        except Exception as e:
            logger.error(f"Error extracting Word content: {str(e)}")
            return "[WORD: Could not process Word document]", {"error": str(e)}

    async def _extract_json_content(self, file_data: bytes) -> Tuple[str, Dict]:
        """Extract and format JSON content."""
        try:
            import json

            text = file_data.decode("utf-8")
            json_data = json.loads(text)

            # Pretty format the JSON for better readability
            formatted_json = json.dumps(json_data, indent=2)

            # Limit text length
            if len(formatted_json) > self.max_text_length:
                formatted_json = (
                    formatted_json[: self.max_text_length] + "... [truncated]"
                )

            return formatted_json, {
                "type": "json",
                "text_length": len(formatted_json),
                "is_valid_json": True,
            }

        except Exception as e:
            logger.error(f"Error extracting JSON: {str(e)}")
            return "[JSON: Could not process JSON file]", {"error": str(e)}


# Global instance
content_extraction_service = ContentExtractionService()
